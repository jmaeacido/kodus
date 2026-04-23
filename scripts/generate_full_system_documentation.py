from __future__ import annotations

import json
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
ENV_FILE = ROOT / ".env"
OUTPUT_DIR = ROOT / "docs"
DOCX_PATH = OUTPUT_DIR / "KODUS_Full_System_Documentation.docx"
PDF_PATH = OUTPUT_DIR / "KODUS_Full_System_Documentation.pdf"


def read_env(path: Path) -> dict[str, str]:
    env: dict[str, str] = {}
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        env[key.strip()] = value.strip().strip('"')
    return env


def connect_db():
    raise NotImplementedError("Database access is delegated to PHP in this environment.")


def fetch_live_db_snapshot() -> dict:
    php_script = r"""
<?php
require __DIR__ . '/../env_helpers.php';
app_load_environment();
$conn = new mysqli(
    app_env('DB_HOST', '127.0.0.1'),
    app_env('DB_USERNAME', 'root'),
    app_env('DB_PASSWORD', ''),
    app_env('DB_NAME', '')
);
if ($conn->connect_error) {
    fwrite(STDERR, $conn->connect_error);
    exit(1);
}
$conn->set_charset('utf8mb4');

$selectedTables = [
    'users','meb','incoming','outgoing','aatracker','breakdown',
    'fund_monitoring_items','fund_monitoring_entries','fund_monitoring_object_codes',
    'project_variable_config','project_lawa_binhi_targets','project_target_entries',
    'program_activity_metadata','program_activity_actual_projects',
    'contact_messages','contact_replies','contact_message_recipients','message_reads',
    'app_notifications','app_notification_reads','events','event_schedule_days',
    'draggable_events','crossmatch_jobs','crossmatch_results',
    'deduplication_jobs','deduplication_results','deduplication_template_outputs',
    'mebis_consolidator_outputs','mebis_lgu_template_outputs',
    'audit_logs','mail_logs','app_settings','meb_change_history'
];

$out = [];
$res = $conn->query("SELECT DATABASE() AS db_name");
$out['database_name'] = ($res && ($row = $res->fetch_assoc())) ? $row['db_name'] : null;

$tableRows = [];
$res = $conn->query("
    SELECT TABLE_NAME, TABLE_ROWS
    FROM information_schema.TABLES
    WHERE TABLE_SCHEMA = DATABASE()
    ORDER BY TABLE_NAME
");
if ($res) {
    while ($row = $res->fetch_assoc()) {
        $tableRows[] = $row;
    }
}
$out['table_rows'] = $tableRows;

$userTypes = [];
$res = $conn->query("
    SELECT userType, COUNT(*) AS count
    FROM users
    GROUP BY userType
    ORDER BY userType
");
if ($res) {
    while ($row = $res->fetch_assoc()) {
        $userTypes[] = $row;
    }
}
$out['user_type_counts'] = $userTypes;

$appSettings = [];
$res = $conn->query("SELECT setting_key, setting_value FROM app_settings ORDER BY setting_key");
if ($res) {
    while ($row = $res->fetch_assoc()) {
        $appSettings[] = $row;
    }
}
$out['app_settings'] = $appSettings;

$columnMap = [];
foreach ($selectedTables as $table) {
    $stmt = $conn->prepare("
        SELECT COLUMN_NAME
        FROM information_schema.COLUMNS
        WHERE TABLE_SCHEMA = DATABASE() AND TABLE_NAME = ?
        ORDER BY ORDINAL_POSITION
    ");
    $stmt->bind_param('s', $table);
    $stmt->execute();
    $result = $stmt->get_result();
    $columns = [];
    while ($row = $result->fetch_assoc()) {
        $columns[] = $row['COLUMN_NAME'];
    }
    if ($columns) {
        $columnMap[$table] = $columns;
    }
    $stmt->close();
}
$out['column_map'] = $columnMap;

echo json_encode($out, JSON_PRETTY_PRINT | JSON_UNESCAPED_SLASHES);
?>
"""
    result = subprocess.run(
        ["php"],
        input=php_script,
        text=True,
        capture_output=True,
        cwd=str(ROOT / "scripts"),
        check=True,
    )
    snapshot = json.loads(result.stdout)
    snapshot["table_count_lookup"] = {
        row["TABLE_NAME"]: int(row["TABLE_ROWS"] or 0)
        for row in snapshot["table_rows"]
    }
    return snapshot


def set_doc_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_sections(snapshot: dict) -> list[dict]:
    table_count_lookup = snapshot["table_count_lookup"]
    user_type_counts = {
        row["userType"]: int(row["count"]) for row in snapshot["user_type_counts"]
    }
    maintenance_settings = {
        row["setting_key"]: row["setting_value"] for row in snapshot["app_settings"]
    }

    data_overview_rows = [
        ("users", "Account identities, profile fields, role, session/security metadata, 2FA data."),
        ("meb", "Masterlist of Eligible Beneficiaries and related classification data."),
        ("incoming", "Incoming document tracking records and uploaded attachment metadata."),
        ("outgoing", "Outgoing document tracking records and forwarding details."),
        ("aatracker", "Administrative Staff activity tracking records."),
        ("breakdown", "Payout breakdown records by province, municipality, and barangay."),
        ("fund_monitoring_items", "Fund monitoring line items per fiscal year."),
        ("fund_monitoring_entries", "Monthly obligations and disbursement values per item."),
        ("project_variable_config", "Fiscal-year project variables such as payout drivers."),
        ("project_lawa_binhi_targets", "Baseline targets and stored target aggregates."),
        ("project_target_entries", "Per-project target rows under a target record."),
        ("program_activity_metadata", "Implementation activity records and summary fields."),
        ("program_activity_actual_projects", "Actual accomplishment projects, map coordinates, and related details."),
        ("contact_messages", "Mailbox threads."),
        ("contact_replies", "Mailbox replies."),
        ("message_reads", "Per-user read and trash state for mailbox items."),
        ("app_notifications", "System-generated transaction notifications."),
        ("app_notification_reads", "Per-user read state for notifications."),
        ("events", "Calendar events."),
        ("event_schedule_days", "Expanded schedule records for multi-day events."),
        ("crossmatch_jobs", "Crossmatching job headers."),
        ("crossmatch_results", "Crossmatching result payloads."),
        ("deduplication_jobs", "Deduplication job headers."),
        ("deduplication_results", "Deduplication grouped results."),
        ("deduplication_template_outputs", "Generated deduplication-ready workbook history."),
        ("mebis_consolidator_outputs", "Generated MEBIS name-matching template history."),
        ("mebis_lgu_template_outputs", "Generated MEB import template history."),
        ("audit_logs", "User and system audit trail entries."),
        ("mail_logs", "Outgoing mail logging records."),
        ("app_settings", "Application-level settings, including maintenance mode state."),
        ("meb_change_history", "Stored before/after beneficiary change history."),
    ]

    database_rows = [
        [table, f"{table_count_lookup.get(table, 0):,}", purpose]
        for table, purpose in data_overview_rows
    ]

    module_rows = [
        ["Public Access and Account Entry", "Root landing, year selection, login, SSO callback, registration, forgot password, reset password, 2FA verification and setup pages."],
        ["Dashboard", "Role-aware home dashboard with quick links, selected-year context, summary metrics, charts, and fund monitoring highlights."],
        ["Operations: Partner-Beneficiaries (MEB)", "Beneficiary masterlist, batch import for administrators, validation for administrators, Excel export, profile file generation, bulk edit/delete for administrators, and batch deletion for administrators."],
        ["Operations: Incoming and Outgoing Tracking", "Track incoming documents, track outgoing documents, attach or reuse files, generate tracking numbers, update records, and view file metadata."],
        ["Operations: Payout", "Province, municipality, and barangay payout breakdown review, update, grouping, and Excel export. Update access is limited to admin and Administrative Staff."],
        ["Operations: Fund Monitoring", "Fiscal-year object codes, fund line items, monthly obligations/disbursement matrix, totals, and variance/utilization views. Data modification is administrator-only."],
        ["Reporting: Implementation Status", "Baseline Targets, Program Activities, Project Location Maps, Project Location Records, LAWA Summary, and BINHI Summary."],
        ["Reporting: Statistical Reports", "Partner-Beneficiaries Profile, Sectoral Data Summary, PWD summary, and PWD sex disaggregation with Excel export."],
        ["Utilities: Crossmatching", "Database-vs-file and file-vs-file beneficiary crossmatching with rule and threshold options, recent jobs, and export."],
        ["Utilities: Deduplication", "Beneficiary deduplication jobs, duplicate grouping results, result export, and deduplication template generation from validated workbooks."],
        ["Utilities: Admin Generators", "MEBIS Name-Matching Template generator and MEB Import Template generator. These generator directories are enforced as administrator-only."],
        ["Workspace and Collaboration", "Mailbox, reply/edit/delete handling, bulk mail actions, unread/trash state, transaction notifications, and account settings."],
        ["Calendar and Scheduling", "Calendar views, draggable template events, private event support, guest handling, and event schedule days."],
        ["Administration", "Users Management, Project Variables, Password Security, Audit Logs, and Maintenance Mode."],
    ]

    role_rows = [
        [
            "Admin",
            "Full system access, including operations, reporting, utilities, inbox, notifications, settings, calendar, all administration pages, all program target/activity management functions, MEB administrative actions, fund monitoring maintenance, user recovery, role changes, password security actions, 2FA reset actions, maintenance mode, and admin-only generators.",
        ],
        [
            "Administrative Staff (`aa`)",
            "Can access the operations workspace. Can open Incoming, Outgoing, MEB, Fund Monitoring (view), Payout, reports, tools, inbox, notifications, settings, and calendar. Can update payout records. Cannot open administrator-only pages, cannot manage program targets/activities, cannot manage project variables, and does not get administrator-only MEB import/validation/bulk-delete controls.",
        ],
        [
            "Implementation Editor (`editor`)",
            "Can access the operations workspace. Can manage Baseline Targets, Program Activities, and Project Variables. Can view reports, tools, inbox, notifications, settings, calendar, and general operations pages. Does not receive administrator-only functions such as user management, maintenance mode, password security, admin-only MEB controls, or admin-only generators.",
        ],
        [
            "User",
            "Blocked from the operations workspace pages explicitly enforced in code: MEB, Incoming, Outgoing, Payout, Fund Monitoring, MEB validation, MEB batch edit, and related operational pages. Can still access dashboard, calendar, implementation-status viewing pages, sectoral/PWD reports, tools, inbox/mail, notifications, and settings. Baseline Targets is hidden because it requires management rights. Beneficiary Profile is hidden from the standard user navigation.",
        ],
    ]

    sections: list[dict] = [
        {
            "heading": "Introduction",
            "paragraphs": [
                (
                    "KODUS, or the KliMalasakit Online Document Updating System, is a web-based PHP and MySQL application used to manage beneficiary records, document tracking, implementation-status reporting, fund monitoring, payout monitoring, internal messaging, and supporting utility workflows. "
                    "This documentation was prepared from an audit of the current application codebase together with the live database schema and active table inventory on April 17, 2026."
                ),
                (
                    "The document is intended for hosting-requirements submission and therefore describes only the functions, roles, data structures, and restrictions that are currently implemented. "
                    "No proposed modules, planned roles, or unimplemented features are included."
                ),
            ],
        },
        {
            "heading": "Objectives",
            "bullets": [
                "Document the current KODUS implementation in a formal, submission-ready format.",
                "Describe the actual functional scope of the system using the source code and live database as the basis of record.",
                "Identify the implemented user roles and distinguish view access from management access where the code does so.",
                "Summarize the current data model, security mechanisms, administrative controls, and hosting considerations.",
            ],
        },
        {
            "heading": "Scope and Limitations",
            "bullets": [
                "This document reflects the application as implemented in the audited code and live database only.",
                "Database descriptions are based on the active database currently configured as "
                f"`{snapshot['database_name']}` rather than on SQL dump files.",
                "Counts shown in the database overview are live table row estimates captured on April 17, 2026.",
                "Where the interface exposes a page but restricts updates through dedicated endpoints, this document states the distinction explicitly.",
                "Legacy or partially retained files are not treated as primary business modules unless they are accessible and functionally connected to the current application.",
            ],
        },
        {
            "heading": "System Overview",
            "paragraphs": [
                (
                    "KODUS operates as a role-aware portal anchored on a selected fiscal year. After authentication, users are routed to a dashboard and a shared navigation shell that changes available actions according to role. "
                    "The system combines operational data entry, reporting, file-based utility tools, account security controls, internal communication, and administration functions in one interface."
                ),
                (
                    "The core application stack is PHP with MySQL, using server-side sessions, Composer-managed libraries, and an AdminLTE-based user interface. "
                    "Excel-oriented processing and exports are implemented through PhpSpreadsheet. Two-factor authentication is implemented using time-based one-time passwords with QR enrollment and recovery codes."
                ),
            ],
        },
        {"heading": "User Roles and Restrictions", "table": {"headers": ["Role", "Current Implemented Access"], "rows": role_rows}},
        {"heading": "Implemented Modules, Pages, and Features", "table": {"headers": ["Module Area", "Implemented Scope"], "rows": module_rows}},
        {
            "heading": "Major Workflows",
            "bullets": [
                "Authentication and session workflow: year selection, login or SSO login, optional 2FA verification, dashboard entry, session timeout handling, and forced logout when roles change or maintenance mode is enabled.",
                "Beneficiary masterlist workflow: administrators can import beneficiary workbooks into MEB, validate records, run batch actions, edit and delete records, delete batches, export Excel files, and generate profile files; non-admin operational users can open the module without the administrator-only controls.",
                "Document tracking workflow: signed-in operational users can create Incoming and Outgoing tracking records, upload permitted file types, generate tracking numbers, and update existing records. Outgoing tracking can reuse a matching file from an incoming record when no new file is uploaded.",
                "Implementation status workflow: target-setting users manage Baseline Targets, implementation-management users manage Program Activities, and all roles with access to the reporting area can review location maps, location records, LAWA summary, and BINHI summary outputs.",
                "Payout workflow: admin and Administrative Staff users open the payout workspace, review grouped payout records, update paid counts and payout dates, and export results. Daily payout values are computed using fiscal-year project variables.",
                "Fund monitoring workflow: administrators maintain object codes, fund line items, and monthly obligations/disbursement entries; viewers can review totals, utilization, variance, and itemized yearly matrices.",
                "Mailbox and notifications workflow: users send and receive mail threads, reply, edit and delete replies where permitted, manage trash state, receive live mail indicators, and review system-generated transaction notifications with per-user read state.",
                "Utility workflow: signed-in users can run Crossmatching and Deduplication jobs and export results; administrators additionally use the MEBIS Name-Matching Template generator and the MEB Import Template generator.",
                "Calendar workflow: users create and manage events, schedule days, draggable event templates, and guest lists, with support for private events in the schema.",
            ],
        },
        {
            "heading": "Database and Data Overview",
            "paragraphs": [
                (
                    f"The live database configured for the current application is `{snapshot['database_name']}`. "
                    f"The active schema currently contains {len(snapshot['table_rows'])} tables. "
                    "User role distribution in the live `users` table at the time of audit is: "
                    f"admin = {user_type_counts.get('admin', 0)}, editor = {user_type_counts.get('editor', 0)}, "
                    f"Administrative Staff (`aa`) = {user_type_counts.get('aa', 0)}, user = {user_type_counts.get('user', 0)}."
                ),
                (
                    "Application-level maintenance settings are stored in `app_settings`. At the time of audit, "
                    f"`maintenance_enabled` = {maintenance_settings.get('maintenance_enabled', '')}, "
                    f"`maintenance_redirect_seconds` = {maintenance_settings.get('maintenance_redirect_seconds', '')}, and "
                    f"`maintenance_warning_seconds` = {maintenance_settings.get('maintenance_warning_seconds', '')}."
                ),
            ],
            "table": {"headers": ["Live Table", "Approx. Live Rows", "Current Purpose"], "rows": database_rows},
        },
        {
            "heading": "Key Data Structures",
            "bullets": [
                "User and access data is centered on `users`, with fields for role (`userType`), profile information, password-policy markers, remember-me and reset tokens, two-factor secrets and recovery codes, last activity, theme preference, and forced role-change/logout state.",
                "Beneficiary data is centered on `meb`, which stores identity fields, location hierarchy, demographic fields, sectoral classifications, validation state, batch identifier, and creation timestamp. `meb_change_history` stores before/after change payloads and edit reasons.",
                "Implementation target data is stored in `project_lawa_binhi_targets` and `project_target_entries`. The live schema supports project-level row IDs, per-project classifications, project types, purok lists, coordinates, latitude/longitude, fertilizer targets, binhi quantities, aquatic resource values, and aggregate LAWA/BINHI target fields.",
                "Implementation activity data is stored in `program_activity_metadata` and `program_activity_actual_projects`. The schema includes province, municipality, barangay, forum dates, stages, liquidation and monitoring dates, partner-beneficiary counts, accomplishment values, land and ownership details, map coordinates, drive links, and actual project relationships.",
                "Document tracking data is stored in `incoming`, `outgoing`, and `aatracker`. These tables store descriptions, dates, generated tracking numbers, file metadata, remarks, receiving or focal details, and user log fields.",
                "Financial monitoring data is distributed across `fund_monitoring_object_codes`, `fund_monitoring_items`, `fund_monitoring_entries`, and `breakdown`, with `project_variable_config` supplying configurable fiscal-year variables that are reused in business rules such as payout rate computation.",
                "Mailbox and notification data is stored in `contact_messages`, `contact_replies`, `contact_message_recipients`, `message_reads`, `app_notifications`, and `app_notification_reads`.",
                "Utility job processing is stored in `crossmatch_jobs`, `crossmatch_results`, `deduplication_jobs`, `deduplication_results`, and the generated-output history tables for deduplication and MEBIS tools.",
                "Scheduling data is stored in `events`, `event_schedule_days`, `draggable_events`, and `event_guests`.",
            ],
        },
        {
            "heading": "Security Features",
            "bullets": [
                "Session hardening is enabled using strict-mode sessions, cookie-only sessions, HTTP-only cookies, same-site cookie configuration, and secure cookies when HTTPS is detected on a non-local host.",
                "CSRF protection is enforced through generated tokens and token checks on modifying requests.",
                "Same-origin enforcement checks request origin, referer, and `Sec-Fetch-Site` for state-changing HTTP methods.",
                "Security headers are applied, including `X-Frame-Options`, `X-Content-Type-Options`, `Referrer-Policy`, `Permissions-Policy`, and HSTS when HTTPS is active.",
                "Password-strength validation requires minimum length plus mixed character classes. Additional password-policy fields track policy version, change dates, and forced password-change state.",
                "Remember-me and reset tokens are hashed and verified through dedicated helpers rather than stored only as raw values.",
                "Two-factor authentication is implemented through TOTP secrets, QR-code enrollment, authenticator verification, recovery-code generation, recovery-code consumption, administrator-triggered reset actions, and setup/disable/regeneration flows.",
                "Role changes and account deactivation can schedule a forced logout so the effective permission state is applied immediately.",
                "Maintenance mode is enforced centrally through application settings and logout/redirect behavior.",
                "Administrative and security-sensitive actions are recorded in `audit_logs`, while outgoing email attempts are recorded in `mail_logs`.",
                "File uploads in the document-tracking modules are type-checked against explicitly allowed extensions and MIME types before being stored.",
            ],
        },
        {
            "heading": "UI and Page Overview",
            "bullets": [
                "The interface uses a common header, sidebar, notification area, and page loader so that role-based navigation remains consistent across modules.",
                "The dashboard is role-aware and changes quick actions and summary cards according to the signed-in role and the selected fiscal year.",
                "Most operational and reporting pages are tabular and rely on DataTables-style interactive listing, sorting, filtering, and export actions.",
                "Implementation Status pages combine tables, modal editing, and a map interface for project coordinates and drive-link references.",
                "Mailbox and settings pages are styled as dedicated workspaces rather than simple forms, and both support light and dark themes.",
                "The calendar module supports full-page event management with filters, agenda views, and draggable event templates.",
                "The utility modules provide upload-driven workflows with recent-job histories and downloadable outputs.",
            ],
        },
        {
            "heading": "Reports and Outputs",
            "bullets": [
                "MEB Excel export.",
                "Partner-Beneficiaries profile file generation from MEB data.",
                "MEB validation export for administrators.",
                "Partner-Beneficiaries Profile Excel export.",
                "Sectoral Data Summary Excel export.",
                "PWD summary Excel export.",
                "PWD sex-disaggregation Excel export.",
                "Payout Excel export.",
                "Crossmatching CSV export.",
                "Deduplication result export.",
                "Program targets Excel template generation and project-target import support.",
                "Deduplication-ready workbook generation from validated MEB workbooks.",
                "MEBIS name-matching template generation.",
                "MEB import template generation.",
                "Recovery code printable output for 2FA setup.",
            ],
        },
        {
            "heading": "Technical and Hosting Notes",
            "bullets": [
                "Application stack: PHP web application with MySQL database and Composer-managed dependencies.",
                "Current Composer packages used by implemented features include `phpoffice/phpspreadsheet`, `phpmailer/phpmailer`, `vlucas/phpdotenv`, `pragmarx/google2fa`, and `bacon/bacon-qr-code`.",
                "Database connection settings are environment-driven through `.env`; the current live database name is "
                f"`{snapshot['database_name']}`.",
                "The application expects writable locations for uploaded and generated files, including utility output folders, generated template folders, cache storage, and document upload directories.",
                "The system depends on PHP sessions, mail configuration for notification/reset messaging, and the configured OAuth endpoints when SSO is enabled.",
                "Character handling is configured for `utf8mb4`, and the application applies MySQL timezone handling from the runtime configuration.",
                "Because multiple business modules generate or store files, the hosting environment should provide sufficient disk space, file-upload limits aligned to operational use, and backup coverage for both database and generated artifacts.",
            ],
        },
        {
            "heading": "Maintenance and Administrative Functions",
            "bullets": [
                "Users Management consolidates user classification, online/idle/offline presence, role changes, deactivation, restoration, and single or bulk 2FA reset actions.",
                "Project Variables provides fiscal-year keyed configuration values used by the application in operational computations and display logic.",
                "Password Security provides administrator actions for password-reset enforcement and reminder/resend operations.",
                "Audit Logs provides searchable visibility into recorded administrative and user actions.",
                "Maintenance Mode provides administrator-controlled warning and redirect behavior stored in `app_settings`.",
                "Soft-delete restoration exists for user accounts, and account role changes/deactivations trigger session handling that prevents stale permissions from persisting.",
            ],
        },
        {
            "heading": "Conclusion",
            "paragraphs": [
                (
                    "KODUS is currently implemented as an integrated web platform that combines beneficiary record administration, document tracking, payout and fund monitoring, implementation-status reporting, collaboration tools, utility processing, and account administration under a single role-aware interface."
                ),
                (
                    "Based on the audited source code and the live database in use on April 17, 2026, the system is production-structured for hosting, with clear separation between administrator functions, implementation-management functions, operational access, and standard-user restrictions. "
                    "This document may therefore be used as the current full-system description for hosting-requirements submission, subject to future updates when the application implementation changes."
                ),
            ],
        },
    ]

    return sections


def configure_docx_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(12.5)


def build_docx(sections: list[dict], snapshot: dict) -> None:
    document = Document()
    configure_docx_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n")
    title = p.add_run("KODUS Full System Documentation")
    title.bold = True
    title.font.size = Pt(22)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("KliMalasakit Online Document Updating System")
    r.font.size = Pt(15)
    r.bold = True

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Submission-Ready Documentation for Hosting Requirements").italic = True

    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(
        "Prepared from the audited application codebase and the live database schema in active use"
    )

    p5 = document.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.add_run("Audit Date: April 17, 2026")

    p6 = document.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.add_run(f"Live Database: {snapshot['database_name']}")

    document.add_page_break()

    for section_data in sections:
        document.add_heading(section_data["heading"], level=1)

        for paragraph in section_data.get("paragraphs", []):
            p = document.add_paragraph(paragraph)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if "bullets" in section_data:
            add_bullets(document, section_data["bullets"])

        if "table" in section_data:
            table_data = section_data["table"]
            rows = table_data["rows"]
            table = document.add_table(rows=len(rows) + 1, cols=len(table_data["headers"]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            for idx, header in enumerate(table_data["headers"]):
                header_cell = table.rows[0].cells[idx]
                set_doc_cell_text(header_cell, header, bold=True)
                shade_cell(header_cell, "D9EAF7")

            for row_index, row_values in enumerate(rows, start=1):
                for col_index, value in enumerate(row_values):
                    set_doc_cell_text(table.rows[row_index].cells[col_index], str(value))

            document.add_paragraph("")

    document.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="KodusTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=28,
            alignment=TA_CENTER,
            spaceAfter=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KodusSubtitle",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KodusBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KodusHeading",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#123B63"),
            spaceBefore=8,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KodusMeta",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KodusBullet",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            alignment=TA_LEFT,
        )
    )
    return styles


def make_pdf_table(headers: list[str], rows: list[list[str]]) -> Table:
    table = Table([headers] + rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#123B63")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#A7BDD6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FBFE")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def add_page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 9)
    canvas.setFillColor(colors.HexColor("#4B5F73"))
    canvas.drawRightString(7.4 * inch, 0.55 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(sections: list[dict], snapshot: dict) -> None:
    styles = pdf_styles()
    story = [
        Spacer(1, 1.5 * inch),
        Paragraph("KODUS Full System Documentation", styles["KodusTitle"]),
        Paragraph("KliMalasakit Online Document Updating System", styles["KodusSubtitle"]),
        Spacer(1, 0.2 * inch),
        Paragraph("Submission-Ready Documentation for Hosting Requirements", styles["KodusMeta"]),
        Paragraph(
            "Prepared from the audited application codebase and the live database schema in active use",
            styles["KodusMeta"],
        ),
        Paragraph("Audit Date: April 17, 2026", styles["KodusMeta"]),
        Paragraph(f"Live Database: {snapshot['database_name']}", styles["KodusMeta"]),
        PageBreak(),
    ]

    for section_data in sections:
        story.append(Paragraph(section_data["heading"], styles["KodusHeading"]))

        for paragraph in section_data.get("paragraphs", []):
            story.append(Paragraph(paragraph.replace("`", ""), styles["KodusBody"]))

        if "bullets" in section_data:
            bullet_items = [
                ListItem(Paragraph(item.replace("`", ""), styles["KodusBullet"]))
                for item in section_data["bullets"]
            ]
            story.append(
                ListFlowable(
                    bullet_items,
                    bulletType="bullet",
                    leftIndent=16,
                    bulletFontName="Helvetica",
                    bulletFontSize=10,
                )
            )
            story.append(Spacer(1, 0.1 * inch))

        if "table" in section_data:
            table_data = section_data["table"]
            story.append(make_pdf_table(table_data["headers"], table_data["rows"]))
            story.append(Spacer(1, 0.12 * inch))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="KODUS Full System Documentation",
        author="OpenAI Codex",
    )
    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    snapshot = fetch_live_db_snapshot()
    sections = build_sections(snapshot)
    build_docx(sections, snapshot)
    build_pdf(sections, snapshot)

    summary = {
        "docx": str(DOCX_PATH),
        "pdf": str(PDF_PATH),
        "database": snapshot["database_name"],
        "tables": len(snapshot["table_rows"]),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
